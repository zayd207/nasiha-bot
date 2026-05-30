from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def generate_doc(data: dict) -> str:
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.color.rgb = RGBColor(15, 27, 51)

    def add_field(label, value):
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(str(value)).font.size = Pt(11)

    def add_divider():
        p = doc.add_paragraph("-" * 55)
        p.runs[0].font.color.rgb = RGBColor(180, 180, 180)
        p.runs[0].font.size = Pt(9)

    # Title
    t = doc.add_heading("NASIHA — BUYURTMA SHAKLI", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = RGBColor(15, 27, 51)

    sub = doc.add_paragraph("Mehr va Tarbiya Olami")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(212, 175, 55)
    sub.runs[0].font.size = Pt(12)

    doc.add_paragraph()
    add_divider()

    # Client
    add_heading("MIJOZ MA'LUMOTLARI", level=1)
    add_field("Ism va Familiya", data.get("client_name", ""))
    add_field("Telefon", data.get("client_phone", ""))
    add_field("Manzil", data.get("client_city", ""))
    add_field("Buyurtma sanasi", datetime.now().strftime("%d.%m.%Y %H:%M"))
    add_field("Bolalar soni", len(data.get("children", [])))
    add_divider()

    # Children
    for i, child in enumerate(data.get("children", []), 1):
        add_heading(f"{i}-BOLA: {child.get('name', '').upper()}", level=1)
        add_field("Ismi", child.get("name", ""))
        add_field("Yoshi", child.get("age", ""))
        add_field("Jinsi", child.get("gender", ""))
        add_field("Xarakteri", child.get("character", ""))
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Muammo tavsifi:").bold = True
        doc.add_paragraph(child.get("problem", ""))
        add_field("Qahramon turi", child.get("hero", ""))
        add_field("Rasmlar soni", f"{len(child.get('photos', []))} ta")
        add_field("Ishtirokchi rasmlari", f"{len(child.get('participant_photos', []))} ta")
        add_divider()

    # Notes
    notes = data.get("extra_notes", "")
    if notes and notes.lower() not in ("yo'q", "yoq", "no", "-"):
        add_heading("QO'SHIMCHA IZOHLAR", level=1)
        doc.add_paragraph(notes)
        add_divider()

    # Footer
    doc.add_paragraph()
    f = doc.add_paragraph("NASIHA — Bola qalbidagi yaxshilik uchun sehrli eshak.")
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f.runs[0].font.color.rgb = RGBColor(212, 175, 55)
    f.runs[0].font.size = Pt(10)

    name_safe = data.get("client_name", "Buyurtma").replace(" ", "_")
    path = f"/tmp/Buyurtma_{name_safe}.docx"
    doc.save(path)
    return path
